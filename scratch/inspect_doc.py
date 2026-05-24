import docx
import os

doc_path = "manual.docx"
if not os.path.exists(doc_path):
    print(f"Error: {doc_path} not found.")
    exit(1)

doc = docx.Document(doc_path)
print(f"Document loaded successfully. Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")

# Print paragraphs that look like placeholders or mentions of screenshots
print("\n--- Paragraphs mentioning images/screenshots ---")
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    # Look for placeholders like [Imagen ...], [Screenshot ...], [Captura ...], or png mentions
    text_lower = text.lower()
    if any(keyword in text_lower for keyword in ["imagen", "screenshot", "captura", ".png", "jpg"]):
        print(f"[{idx}]: {text}")

# Check for inline shapes (images already in document)
print(f"\nInline shapes: {len(doc.inline_shapes)}")
for idx, shape in enumerate(doc.inline_shapes):
    print(f"Shape {idx}: type={shape.type}, width={shape.width}, height={shape.height}")

# Let's inspect some content or outline of the document (e.g. headings)
print("\n--- Document Outline (Headings) ---")
for idx, p in enumerate(doc.paragraphs):
    if p.style.name.startswith("Heading"):
        print(f"[{idx}] {p.style.name}: {p.text}")
