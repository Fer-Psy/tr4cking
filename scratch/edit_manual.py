"""
Script para agregar la sección de Reportes al manual.docx de TR4CKING.
Agrega al final del documento una sección con capturas de pantalla 
de los reportes más importantes y sus funcionalidades.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

MANUAL_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'manual.docx')
SCREENSHOTS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'screenshots')

def add_section_heading(doc, text, level=1):
    """Add a heading with custom formatting."""
    heading = doc.add_heading(text, level=level)
    return heading

def add_figure_caption(doc, text):
    """Add a centered figure caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    return p

def add_body_text(doc, text):
    """Add a normal paragraph."""
    p = doc.add_paragraph(text)
    return p

def add_bullet_list(doc, items):
    """Add a bullet list."""
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def add_image_safe(doc, img_path, width=Inches(6.0)):
    """Add an image centered, handling errors gracefully."""
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=width)
        return True
    else:
        print(f"WARNING: Image not found: {img_path}")
        return False

def main():
    print(f"Opening manual: {MANUAL_PATH}")
    doc = Document(MANUAL_PATH)
    
    # =========================================================================
    # Add page break before the new section
    # =========================================================================
    doc.add_page_break()
    
    # =========================================================================
    # MAIN SECTION: REPORTES DEL SISTEMA
    # =========================================================================
    add_section_heading(doc, "REPORTES DEL SISTEMA", level=1)
    
    add_body_text(doc, 
        "El sistema TR4CKING cuenta con un completo modulo de reportes que permite a los administradores "
        "y personal autorizado obtener informacion detallada sobre las operaciones de la empresa. "
        "Los reportes se acceden desde el menu lateral en la seccion FINANZAS > Centro de Reportes, "
        "y desde la seccion de Caja. A continuacion se detallan los reportes mas importantes del sistema "
        "con sus funcionalidades."
    )
    
    # =========================================================================
    # 1. CENTRO DE REPORTES (REPORTE DIARIO)
    # =========================================================================
    doc.add_page_break()
    add_section_heading(doc, "Centro de Reportes", level=2)
    
    add_body_text(doc,
        "El Centro de Reportes es la herramienta principal de analisis del sistema. Permite consultar "
        "datos operativos y financieros dentro de un rango de fechas y filtrar por empresa. "
        "Cuenta con las siguientes funcionalidades generales:"
    )
    
    add_bullet_list(doc, [
        "Filtro por rango de fechas (Fecha Desde / Fecha Hasta)",
        "Filtro por empresa (o seleccionar todas las empresas)",
        "Boton Consultar para ejecutar la busqueda",
        "Boton Imprimir para imprimir el reporte directamente desde el navegador",
        "Boton Exportar PDF para descargar el reporte en formato PDF",
        "Navegacion por pestanas para acceder a los diferentes tipos de reportes",
    ])
    
    add_body_text(doc,
        "El Centro de Reportes se organiza en 6 pestanas principales: Flujo de Ingresos, "
        "Cierres de Caja, Encomiendas, Situacion Diaria, Comisiones (solo administradores) y Ventas."
    )
    
    # Screenshot del Centro de Reportes
    img_path = os.path.join(SCREENSHOTS_DIR, '33_reporte_diario.png')
    add_image_safe(doc, img_path, width=Inches(6.2))
    add_figure_caption(doc, "Figura 46: Centro de Reportes - Vista general con pestanas de navegacion")
    
    # =========================================================================
    # 1.1 FLUJO DE INGRESOS
    # =========================================================================
    add_section_heading(doc, "Reporte: Flujo de Ingresos", level=3)
    
    add_body_text(doc,
        "La pestana de Flujo de Ingresos presenta un resumen ejecutivo de los ingresos operativos. "
        "Es la vista predeterminada al ingresar al Centro de Reportes."
    )
    
    add_body_text(doc, "Funcionalidades principales:")
    
    add_bullet_list(doc, [
        "Tarjetas de estadisticas rapidas: total de viajes realizados, pasajes vendidos, encomiendas "
        "registradas e ingresos totales en guaranies",
        "Desglose de ingresos por pasajes y por encomiendas en tarjetas separadas",
        "Grafico de barras interactivo que muestra los ingresos por dia, diferenciando pasajes "
        "y encomiendas con colores distintos",
        "Tabla detallada de viajes del periodo con: ID, fecha, itinerario, bus, chofer, cantidad "
        "de pasajes, porcentaje de ocupacion (con barra de progreso visual) y estado del viaje",
    ])
    
    # =========================================================================
    # 1.2 CIERRES DE CAJA
    # =========================================================================
    add_section_heading(doc, "Reporte: Cierres de Caja", level=3)
    
    add_body_text(doc,
        "La pestana de Cierres de Caja permite revisar todas las sesiones de caja abiertas "
        "y cerradas en el periodo seleccionado."
    )
    
    add_body_text(doc, "Funcionalidades principales:")
    
    add_bullet_list(doc, [
        "Tabla de sesiones de caja con: cajero, fecha/hora de apertura y cierre, duracion de la sesion, "
        "monto de apertura, total de ingresos, total de egresos, monto de cierre real, "
        "diferencia (sobrante/faltante) y estado",
        "Indicadores de diferencia con codigos de color: verde (sin diferencia), amarillo (sobrante), "
        "rojo (faltante)",
        "Grafico de barras comparativo que muestra apertura, ingresos y egresos por cada sesion de caja",
    ])
    
    # =========================================================================
    # 1.3 ENCOMIENDAS
    # =========================================================================
    add_section_heading(doc, "Reporte: Encomiendas", level=3)
    
    add_body_text(doc,
        "La pestana de Encomiendas presenta el detalle de todas las encomiendas del periodo."
    )
    
    add_body_text(doc, "Funcionalidades principales:")
    
    add_bullet_list(doc, [
        "Tabla con informacion de cada encomienda: codigo, fecha, remitente, destinatario, "
        "origen, destino, precio y estado",
        "Estadisticas por estado de encomienda (registrado, en transito, en destino, entregado, cancelado) "
        "con conteo y montos",
    ])
    
    # =========================================================================
    # 1.4 SITUACION DIARIA
    # =========================================================================
    add_section_heading(doc, "Reporte: Situacion Diaria", level=3)
    
    add_body_text(doc,
        "La pestana de Situacion Diaria ofrece un panorama completo de pasajes y encomiendas "
        "del periodo seleccionado, combinando tarjetas de resumen, graficos tipo donut y tablas detalladas."
    )
    
    add_body_text(doc, "Funcionalidades principales:")
    
    add_bullet_list(doc, [
        "Resumen de Pasajes por estado: cantidad y monto total para cada estado "
        "(vendido, reservado, abordado, cancelado)",
        "Resumen de Encomiendas por estado: cantidad y monto total para cada estado "
        "(registrado, en transito, en destino, entregado, cancelado)",
        "Graficos tipo donut para visualizar la distribucion de estados de pasajes y encomiendas",
        "Tabla detallada de pasajes: codigo, fecha, pasajero, ruta (origen a destino), asiento, "
        "vendedor, precio y estado",
        "Tabla detallada de encomiendas: codigo, fecha, remitente, destinatario, destino, "
        "registrador, precio y estado",
    ])
    
    # =========================================================================
    # 1.5 COMISIONES
    # =========================================================================
    add_section_heading(doc, "Reporte: Comisiones Mensuales", level=3)
    
    add_body_text(doc,
        "La pestana de Comisiones (disponible solo para administradores) muestra el calculo de "
        "comisiones por agente comercial en el periodo seleccionado. Permite controlar las comisiones "
        "devengadas por cada agente sobre las ventas de pasajes y encomiendas."
    )
    
    # =========================================================================
    # 1.6 VENTAS
    # =========================================================================
    add_section_heading(doc, "Reporte: Ventas", level=3)
    
    add_body_text(doc,
        "La pestana de Ventas presenta un reporte consolidado de todas las ventas del periodo."
    )
    
    # Screenshot del Reporte de Ventas
    img_path = os.path.join(SCREENSHOTS_DIR, '34_reporte_ventas.png')
    add_image_safe(doc, img_path, width=Inches(6.2))
    add_figure_caption(doc, "Figura 47: Centro de Reportes - Pestana de Ventas")
    
    add_body_text(doc, "Funcionalidades principales:")
    
    add_bullet_list(doc, [
        "Tarjetas de resumen: total pasajes (cantidad e ingresos), total encomiendas (cantidad e ingresos), "
        "facturas emitidas e ingresos totales",
        "Tabla de ventas por dia con desglose de: fecha, cantidad de pasajes, cantidad de encomiendas, "
        "facturas, ingresos por pasajes, ingresos por encomiendas y total diario",
        "Fila de totales al final de la tabla con los acumulados del periodo",
        "Resumen de IVA con bases imponibles al 5% y 10%, montos exentos y liquidacion de impuestos",
        "Grafico de distribucion de ingresos con barras de progreso que muestran el porcentaje "
        "de participacion de pasajes y encomiendas sobre el total",
    ])
    
    # =========================================================================
    # 2. REPORTE DE VENTAS (VISTA SEPARADA)
    # =========================================================================
    doc.add_page_break()
    add_section_heading(doc, "Reporte de Ventas (Acceso Directo)", level=2)
    
    add_body_text(doc,
        "Ademas del Centro de Reportes, el sistema cuenta con una vista de Reporte de Ventas "
        "de acceso directo que permite generar informes rapidos de ventas con filtros por tipo "
        "(pasajes, encomiendas o todos) y rango de fechas. Esta vista complementaria ofrece "
        "una interfaz simplificada con las mismas funcionalidades de ventas."
    )
    
    add_body_text(doc, "Funcionalidades:")
    
    add_bullet_list(doc, [
        "Selector de rango de fechas (Fecha Inicio / Fecha Fin)",
        "Filtro por tipo de operacion: todos, pasajes o encomiendas",
        "Boton Generar para ejecutar la consulta",
        "Boton Imprimir para imprimir el reporte",
        "Resumen de totales: pasajes (con ingresos), encomiendas (con ingresos), facturas emitidas "
        "e ingresos totales",
        "Tabla de ventas por dia con desglose completo",
        "Resumen de IVA (base imponible 5%, 10%, exenta, IVA 5%, IVA 10%, total IVA)",
        "Distribucion de ingresos con barras de progreso porcentuales",
    ])
    
    # =========================================================================
    # 3. GESTION DE CAJA
    # =========================================================================
    doc.add_page_break()
    add_section_heading(doc, "Dashboard de Caja", level=2)
    
    add_body_text(doc,
        "El Dashboard de Caja es una vista operativa que permite al cajero gestionar su sesion "
        "de caja en tiempo real, registrando todos los movimientos financieros del dia."
    )
    
    # Screenshot del Dashboard de Caja
    img_path = os.path.join(SCREENSHOTS_DIR, '32_caja_dashboard.png')
    add_image_safe(doc, img_path, width=Inches(6.2))
    add_figure_caption(doc, "Figura 48: Dashboard de Caja con sesion abierta")
    
    add_body_text(doc, "Funcionalidades del Dashboard de Caja:")
    
    add_bullet_list(doc, [
        "Indicador visual del estado de la caja (ABIERTA en verde / CERRADA en gris) con hora de apertura",
        "Resumen de caja en tiempo real: monto de apertura, total de ingresos (en verde), "
        "total de egresos (en rojo) y total actual calculado automaticamente",
        "Tabla de ultimos movimientos con columnas: hora, empresa, tipo (ingreso/egreso con badge de color), "
        "concepto, detalle de operacion (con boton expandible para ver items individuales) y monto",
        "Los detalles de operacion muestran informacion especifica: para pasajes el numero de venta, "
        "asiento y nombre del pasajero; para encomiendas el numero de encomienda y remitente",
        "Boton para registrar nuevo movimiento manual (ingreso o egreso no automatico)",
        "Boton para cerrar caja al final del turno",
        "Boton para imprimir informe de caja (genera un reporte detallado optimizado para impresion)",
        "Historial de sesiones anteriores con: fecha, cajero, horario, duracion y diferencia de cierre",
    ])
    
    add_body_text(doc,
        "Al imprimir el informe de caja, el sistema genera automaticamente un reporte con "
        "encabezado TR4CKING, los datos del cajero, las fechas de apertura y cierre, un cuadro "
        "resumen con montos, y la tabla completa de movimientos con numeracion secuencial."
    )
    
    # =========================================================================
    # 4. CLIENTES PENDIENTES DE FACTURAR
    # =========================================================================
    doc.add_page_break()
    add_section_heading(doc, "Reporte de Clientes Pendientes de Facturar", level=2)
    
    add_body_text(doc,
        "Este reporte muestra todos los clientes que tienen pasajes y/o encomiendas sin facturar, "
        "facilitando el proceso de facturacion masiva."
    )
    
    # Screenshot de Pendientes de Facturar
    img_path = os.path.join(SCREENSHOTS_DIR, '36_pendientes_factura.png')
    add_image_safe(doc, img_path, width=Inches(6.2))
    add_figure_caption(doc, "Figura 49: Reporte de Clientes Pendientes de Facturar")
    
    add_body_text(doc, "Funcionalidades del reporte:")
    
    add_bullet_list(doc, [
        "Tarjetas de estadisticas: total de clientes pendientes, pasajes sin facturar, "
        "encomiendas sin facturar y proximos numeros de factura por empresa (Guairena/Ybyturuzu)",
        "Informacion de timbrados vigentes por empresa (numero de timbrado, punto de expedicion "
        "y fecha de vigencia)",
        "Buscador para filtrar por nombre de cliente, pasajero, cedula o codigo de pasaje",
        "Tabla de clientes con: nombre, cedula/RUC, cantidad de pasajes pendientes, "
        "cantidad de encomiendas pendientes, total pasajes, total encomiendas, total general "
        "y acciones (Facturar / Cancelar todo)",
        "Detalle colapsable por cliente que muestra el listado individual de pasajes "
        "(pasajero, ruta, fecha, asiento, precio) y encomiendas (codigo, destino, precio)",
        "Boton directo para facturar que redirige al formulario de factura con los datos "
        "del cliente pre-cargados",
        "Boton para cancelar todos los items pendientes de un cliente",
        "Enlace a la lista de facturas emitidas",
    ])
    
    # =========================================================================
    # 5. FUNCIONALIDADES DE EXPORTACION
    # =========================================================================
    doc.add_page_break()
    add_section_heading(doc, "Funcionalidades de Exportacion e Impresion", level=2)
    
    add_body_text(doc,
        "Todos los reportes del sistema cuentan con funcionalidades de exportacion e impresion "
        "que permiten generar documentos para uso externo."
    )
    
    add_body_text(doc, "Opciones de exportacion disponibles:")
    
    add_bullet_list(doc, [
        "Impresion directa desde el navegador: al hacer clic en el boton Imprimir, el sistema "
        "genera una version optimizada del reporte con encabezado corporativo TR4CKING, datos "
        "del periodo, empresa y usuario, formato blanco y negro para ahorro de tinta, tablas "
        "compactas y un pie de pagina con la fecha de impresion",
        "Exportacion a PDF: el boton Exportar PDF genera un archivo PDF descargable con el mismo "
        "formato optimizado de impresion, en tamano A4 horizontal, ideal para archivar o enviar "
        "por correo electronico",
        "Informe de caja imprimible: desde el Dashboard de Caja, el boton Imprimir Informe genera "
        "un reporte completo de la sesion de caja con todos los movimientos detallados, "
        "en formato A4 vertical",
    ])
    
    add_body_text(doc,
        "Los reportes impresos y exportados incluyen automaticamente un encabezado con el logo "
        "TR4CKING, el tipo de reporte, el periodo consultado, la empresa filtrada, el usuario "
        "que genera el reporte y la fecha/hora de generacion."
    )
    
    # =========================================================================
    # Save the document
    # =========================================================================
    print(f"Saving manual: {MANUAL_PATH}")
    doc.save(MANUAL_PATH)
    print("Manual updated successfully!")
    
    # Verify by counting paragraphs
    doc2 = Document(MANUAL_PATH)
    count = len(doc2.paragraphs)
    print(f"Total paragraphs in updated manual: {count}")

if __name__ == '__main__':
    main()
