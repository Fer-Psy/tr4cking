import docx

doc = docx.Document("manual.docx")

def run_has_shape(run):
    return len(run._r.xpath('.//w:drawing')) > 0

shape_idx = 0
mappings = []

for p_idx, p in enumerate(doc.paragraphs):
    p_shapes = []
    for r_idx, run in enumerate(p.runs):
        if run_has_shape(run):
            p_shapes.append(shape_idx)
            shape_idx += 1
            
    if p_shapes:
        # Find the figure caption in the surrounding paragraphs (within 2 paragraphs after or 2 paragraphs before)
        caption = "Unknown"
        # Search next 3 paragraphs
        for offset in [1, 2, 3]:
            if p_idx + offset < len(doc.paragraphs):
                text = doc.paragraphs[p_idx + offset].text.strip()
                if "Figura" in text:
                    caption = text
                    break
        # If not found, search previous 3 paragraphs
        if caption == "Unknown":
            for offset in [-1, -2, -3]:
                if p_idx + offset >= 0:
                    text = doc.paragraphs[p_idx + offset].text.strip()
                    if "Figura" in text:
                        caption = text
                        break
        
        for s in p_shapes:
            mappings.append((s, p_idx, caption))

print(f"Total shapes mapped: {len(mappings)}")
for m in mappings:
    print(f"Shape {m[0]} (P[{m[1]}]): {m[2]}")
