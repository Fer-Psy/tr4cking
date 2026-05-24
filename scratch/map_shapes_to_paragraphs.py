import docx

doc = docx.Document("manual.docx")

print("Mapping of shapes to paragraphs and surroundings:")
shape_count = 0

# Helper to check if a run has a drawing / shape
def run_has_shape(run):
    return len(run._r.xpath('.//w:drawing')) > 0

for p_idx, p in enumerate(doc.paragraphs):
    p_shapes = []
    for r_idx, run in enumerate(p.runs):
        if run_has_shape(run):
            # Find the shapes in this run
            run_drawings = run._r.xpath('.//w:drawing')
            for d in run_drawings:
                p_shapes.append(shape_count)
                shape_count += 1
                
    if p_shapes:
        # Print surrounding text
        print(f"\n--- Paragraph P[{p_idx}] contains Shapes {p_shapes} ---")
        # Print preceding 2 non-empty paragraphs
        preceding = []
        for i in range(p_idx - 1, -1, -1):
            txt = doc.paragraphs[i].text.strip()
            if txt:
                preceding.append(f"P[{i}]: {txt}")
                if len(preceding) == 2:
                    break
        for txt in reversed(preceding):
            print(f"  [Preceding] {txt}")
            
        print(f"  [Self] P[{p_idx}]: '{p.text.strip()}' (style={p.style.name})")
        
        # Print succeeding 2 non-empty paragraphs
        succeeding = []
        for i in range(p_idx + 1, len(doc.paragraphs)):
            txt = doc.paragraphs[i].text.strip()
            if txt:
                succeeding.append(f"P[{i}]: {txt}")
                if len(succeeding) == 2:
                    break
        for txt in succeeding:
            print(f"  [Succeeding] {txt}")

print(f"\nTotal mapped shapes in paragraphs: {shape_count}")
